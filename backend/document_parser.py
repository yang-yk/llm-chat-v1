"""
文档解析服务
支持 txt、pdf、doc、docx 等格式的文档解析
"""
import os
import tempfile
from pathlib import Path


class DocumentParser:
    """文档解析器"""

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """解析TXT文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            content = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise ValueError(f"无法解码文件 {file_path}，尝试了以下编码: {encodings}")

            return content.strip()
        except Exception as e:
            raise Exception(f"解析TXT文件失败: {str(e)}")

    @staticmethod
    def clean_pdf_text(text: str) -> str:
        """
        清理PDF提取的文本，处理不必要的换行

        策略：
        1. 如果一行以标点符号结尾（。！？；：）或行尾是完整句子，保留换行
        2. 如果一行不以标点结尾，说明是段落内换行，替换为空格
        3. 多个连续空格合并为一个
        """
        import re

        lines = text.split('\n')
        cleaned_lines = []

        # 中英文句尾标点
        sentence_end_marks = ('。', '！', '？', '；', '.', '!', '?', ';', ':', '：')

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                # 空行保留（作为段落分隔）
                if cleaned_lines and cleaned_lines[-1] != '\n':
                    cleaned_lines.append('\n')
                continue

            # 如果当前行以标点结尾，或者是最后一行，作为独立行
            if line.endswith(sentence_end_marks) or i == len(lines) - 1:
                cleaned_lines.append(line + '\n')
            else:
                # 否则，将其与下一行合并（添加空格）
                cleaned_lines.append(line + ' ')

        # 合并所有行
        result = ''.join(cleaned_lines)

        # 处理多余的空格
        result = re.sub(r' +', ' ', result)  # 多个空格合并为一个
        result = re.sub(r' ([。！？；：，、])', r'\1', result)  # 标点前的空格移除
        result = re.sub(r'\n\n\n+', '\n\n', result)  # 多个空行合并为两个

        return result.strip()

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """解析PDF文件"""
        try:
            import PyPDF2

            print(f"[INFO] 开始解析PDF文件: {file_path}")
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                print(f"[INFO] PDF共有 {num_pages} 页")

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            # 合并所有页面的文本
            content = '\n\n'.join(text_content)

            if not content.strip():
                print(f"[WARNING] PDF文件提取的文本为空: {file_path}")
                raise Exception("PDF文件内容为空或无法提取文本，可能是扫描版PDF")

            # 清理文本中不必要的换行
            content = DocumentParser.clean_pdf_text(content)

            print(f"[INFO] PDF解析成功，提取文本长度: {len(content)}")
            return content
        except ImportError:
            raise Exception("需要安装 PyPDF2 库: pip install PyPDF2")
        except Exception as e:
            print(f"[ERROR] 解析PDF失败: {str(e)}")
            raise Exception(f"解析PDF文件失败: {str(e)}")

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """解析DOCX文件"""
        try:
            from docx import Document
            import zipfile

            print(f"[INFO] 开始解析DOCX文件: {file_path}")

            # 验证文件存在
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            # 验证文件大小
            file_size = os.path.getsize(file_path)
            print(f"[INFO] DOCX文件大小: {file_size} 字节")

            if file_size == 0:
                raise Exception("DOCX文件为空")

            # 验证是否是有效的 ZIP 文件（DOCX 本质上是 ZIP）
            if not zipfile.is_zipfile(file_path):
                # 尝试检测是否是旧版 DOC 格式
                with open(file_path, 'rb') as f:
                    header = f.read(8)
                    # DOC 文件通常以 D0 CF 11 E0 开头（OLE2 格式）
                    if header[:4] == b'\xD0\xCF\x11\xE0':
                        raise Exception(
                            f"该文件实际上是旧版 DOC 格式（Word 2003 或更早版本），而不是 DOCX 格式。\n\n"
                            f"解决方案：\n"
                            f"1. 使用 Word 打开文件，另存为 .docx 格式后重新上传\n"
                            f"2. 或将文件扩展名改为 .doc 后上传（系统支持 DOC 格式）"
                        )
                    else:
                        raise Exception(
                            f"文件不是有效的 DOCX 格式。DOCX 文件本质上是 ZIP 压缩包，但该文件不是 ZIP 格式。\n\n"
                            f"可能的原因：\n"
                            f"1. 文件损坏或不完整\n"
                            f"2. 文件是其他格式但扩展名被修改为 .docx\n"
                            f"3. 文件是旧版 DOC 格式\n\n"
                            f"建议：使用 Microsoft Word 重新保存为 .docx 格式后再上传"
                        )

            # 尝试打开 DOCX 文件
            print(f"[INFO] 尝试打开DOCX文件...")
            doc = Document(file_path)
            print(f"[INFO] DOCX文件打开成功")

            text_content = []

            # 提取段落文本
            print(f"[INFO] 提取段落文本...")
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # 提取表格文本
            print(f"[INFO] 提取表格文本...")
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_content.append(row_text)

            content = '\n\n'.join(text_content)
            print(f"[INFO] DOCX解析成功，提取文本长度: {len(content)}")
            return content.strip()

        except ImportError:
            raise Exception("需要安装 python-docx 库: pip install python-docx")
        except FileNotFoundError as e:
            print(f"[ERROR] 文件不存在: {str(e)}")
            raise Exception(f"文件不存在: {str(e)}")
        except Exception as e:
            print(f"[ERROR] 解析DOCX失败: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"解析DOCX文件失败: {str(e)}")

    @staticmethod
    def parse_doc(file_path: str) -> str:
        """解析DOC文件（旧版Word格式）"""
        try:
            import subprocess
            import platform

            print(f"[INFO] 开始解析DOC文件: {file_path}")

            # 方法1: 尝试使用 antiword 工具（Linux，最可靠，支持中文）
            if platform.system() == "Linux":
                try:
                    print(f"[INFO] 尝试使用 antiword 解析DOC文件")
                    result = subprocess.run(
                        ['antiword', '-m', 'UTF-8.txt', file_path],
                        capture_output=True,
                        text=True,
                        check=True,
                        timeout=30
                    )
                    content = result.stdout.strip()
                    if content:
                        print(f"[INFO] antiword 解析成功，文本长度: {len(content)}")
                        return content
                    else:
                        print(f"[WARNING] antiword 解析结果为空")
                except subprocess.TimeoutExpired:
                    print(f"[ERROR] antiword 解析超时")
                except (subprocess.CalledProcessError, FileNotFoundError) as e:
                    print(f"[WARNING] antiword 未安装或解析失败: {str(e)}")

            # 方法2: 尝试使用 catdoc 工具（Linux，另一个选择）
            if platform.system() == "Linux":
                try:
                    print(f"[INFO] 尝试使用 catdoc 解析DOC文件")
                    result = subprocess.run(
                        ['catdoc', '-d', 'utf-8', file_path],
                        capture_output=True,
                        text=True,
                        check=True,
                        timeout=30
                    )
                    content = result.stdout.strip()
                    if content:
                        print(f"[INFO] catdoc 解析成功，文本长度: {len(content)}")
                        return content
                    else:
                        print(f"[WARNING] catdoc 解析结果为空")
                except subprocess.TimeoutExpired:
                    print(f"[ERROR] catdoc 解析超时")
                except (subprocess.CalledProcessError, FileNotFoundError) as e:
                    print(f"[WARNING] catdoc 未安装或解析失败: {str(e)}")

            # 如果都失败，返回提示信息
            raise Exception(
                "无法解析DOC文件（旧版Word格式）。建议：\n\n"
                "方案1（推荐）：将DOC文件另存为DOCX格式后重新上传\n"
                "方案2：在Linux系统上安装解析工具：\n"
                "  - sudo apt-get install antiword\n"
                "  - 或 sudo apt-get install catdoc\n\n"
                "注意：DOC是微软旧版Word格式，解析较为复杂，建议使用DOCX格式"
            )
        except Exception as e:
            print(f"[ERROR] 解析DOC文件最终失败: {str(e)}")
            raise Exception(f"解析DOC文件失败: {str(e)}")

    @classmethod
    def parse_file(cls, file_path: str, file_type: str) -> str:
        """
        根据文件类型解析文件

        Args:
            file_path: 文件路径
            file_type: 文件类型 (txt/pdf/doc/docx)

        Returns:
            解析后的文本内容
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_type = file_type.lower()

        if file_type == 'txt':
            return cls.parse_txt(file_path)
        elif file_type == 'pdf':
            return cls.parse_pdf(file_path)
        elif file_type == 'docx':
            return cls.parse_docx(file_path)
        elif file_type == 'doc':
            return cls.parse_doc(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

    @staticmethod
    def get_file_type(filename: str) -> str:
        """从文件名获取文件类型"""
        ext = Path(filename).suffix.lower()
        if ext.startswith('.'):
            ext = ext[1:]

        supported_types = ['txt', 'pdf', 'doc', 'docx']
        if ext in supported_types:
            return ext
        else:
            raise ValueError(f"不支持的文件类型: {ext}")


# 创建全局实例
document_parser = DocumentParser()
